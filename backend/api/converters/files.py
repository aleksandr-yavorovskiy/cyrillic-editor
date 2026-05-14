from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader

from api.core.base import BaseImporter, BaseExporter
from api.core.exceptions import FileProcessingError, UnsupportedFormatError
from api.core.options import CompileOptions


class TextImporter(BaseImporter):
    def convert(self, source: BytesIO) -> str:
        try:
            return source.read().decode("utf-8")
        except UnicodeDecodeError as e:
            raise FileProcessingError(f"Failed to decode text file: {e}")
        except Exception as e:
            raise FileProcessingError(f"Failed to read text file: {e}")


class DocxImporter(BaseImporter):
    def convert(self, source: BytesIO) -> str:
        try:
            doc = Document(source)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise FileProcessingError(f"Failed to read DOCX file: {e}")


class PdfImporter(BaseImporter):
    def convert(self, source: BytesIO) -> str:
        try:
            reader = PdfReader(source)
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            if not pages:
                raise FileProcessingError("PDF contains no extractable text")
            return "\n".join(pages)
        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"Failed to read PDF file: {e}")


class DocxExporter(BaseExporter):
    def export(self, text: str, options=None) -> BytesIO:
        try:
            document = Document()
            for line in text.split("\n"):
                document.add_paragraph(line)
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return buffer
        except Exception as e:
            raise FileProcessingError(f"Failed to create DOCX file: {e}")


class PdfExporter(BaseExporter):
    def __init__(self, compilation_service=None):
        self._service = compilation_service

    def export(self, text: str, options=None) -> BytesIO:
        from api.services import CompilationService  # lazy import to break circular dep

        opts = options or CompileOptions()
        service = self._service or CompilationService()
        return service.compile(text, opts)


class ConverterRegistry:
    def __init__(self):
        self._importers = {}
        self._exporters = {}
        self._initialize()

    def _initialize(self):
        self.register_importer(".txt", TextImporter())
        self.register_importer(".docx", DocxImporter())
        self.register_importer(".pdf", PdfImporter())
        self.register_exporter(".docx", DocxExporter())
        self.register_exporter(".pdf", PdfExporter())

    def register_importer(self, ext: str, importer: BaseImporter):
        self._importers[ext] = importer

    def register_exporter(self, ext: str, exporter: BaseExporter):
        self._exporters[ext] = exporter

    def get_importer(self, ext: str) -> BaseImporter | None:
        return self._importers.get(ext)

    def get_exporter(self, ext: str) -> BaseExporter | None:
        return self._exporters.get(ext)

    @property
    def importers(self):
        return self._importers

    @property
    def exporters(self):
        return self._exporters

    def detect_extension(self, filename: str) -> str:
        for ext in self._importers.keys():
            if filename.lower().endswith(ext):
                return ext
        raise UnsupportedFormatError(f"Unsupported file format: {filename}")
