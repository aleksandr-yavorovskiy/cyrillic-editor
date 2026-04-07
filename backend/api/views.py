from django.http import JsonResponse, HttpResponse
from rest_framework.decorators import api_view
from pathlib import Path
from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO
import os
import subprocess
import tempfile
from django.conf import settings
import logging


logger = logging.getLogger("api")


# TODO: unify escape_latex & insert_soft_breaks (get O(n))
def escape_latex(text):
    replacements = {
        "&": "\\&",
        "%": "\\%",
        "$": "\\$",
        "#": "\\#",
        "_": "\\_",
        "{": "\\{",
        "}": "\\}",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    return text.replace("\n", "\\\\\n")


def insert_soft_breaks(text, threshold=15):
    words = text.split(" ")
    processed_words = []

    for word in words:
        if len(word) > threshold:
            processed_words.append(f"\\seqsplit{{{word}}}")
        else:
            processed_words.append(word)

    return " ".join(processed_words)


def to_measure_units(margin_num, measure_units, default=2):
    try:
        return f"{margin_num}{measure_units}"
    except:
        return f"{default}{measure_units}"


# TODO: move inner logic to another layer
@api_view(['POST'])
def compile_handler(request):
    text = request.data.get("text", "")
    processed_text = insert_soft_breaks(escape_latex(text))

    font = request.data.get("font", "")
    font_path = settings.BASE_DIR / "fonts"
    font_path_str = str(font_path) + "/"

    font_path = "/home/ysd/spbu-diploma/cyrillic-editor/backend/fonts/"

    # TODO: move units to const/enum
    top = to_measure_units(request.data.get("top"), "cm")
    bottom = to_measure_units(request.data.get("bottom"), "cm")
    left = to_measure_units(request.data.get("left"), "cm")
    right = to_measure_units(request.data.get("right"), "cm")

    font_size = int(request.data.get("fontSize", 14))
    line_height = int(font_size * 1.2)

    tex_content = f"""
\\documentclass{{article}}
\\usepackage{{fontspec}}
\\usepackage{{seqsplit}}
\\usepackage[margin=1in]{{geometry}}

\\setmainfont{{{font}}}[
    Path={font_path_str}
]

\\geometry{{
  top={top},
  bottom={bottom},
  left={left},
  right={right}
}}

\\pagestyle{{empty}}

\\begin{{document}}
\\noindent {{\\fontsize{{{font_size}}}{{{line_height}}}\\selectfont {processed_text}}}\\end{{document}}
"""

    tex_file = "file.tex"
    pdf_file = "file.pdf"

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_path = os.path.join(tmpdir, tex_file)

            with open(tex_path, "w", encoding="utf-8") as f:
                f.write(tex_content)

            subprocess.run(
                ["xelatex", "-interaction=nonstopmode", tex_file],
                cwd=tmpdir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

            pdf_path = os.path.join(tmpdir, pdf_file)

            if not os.path.exists(pdf_path):
                return HttpResponse("PDF not generated", status=500)

            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

            response = HttpResponse(pdf_bytes, content_type="application/pdf")
            response["Content-Disposition"] = "inline; filename=result.pdf"
            return response

    except Exception as e:
        logger.debug(f"Couldn't compile. Error: {str(e)}")
        return HttpResponse(f"Error: {str(e)}", status=500)


@api_view(['GET'])
def ping(request):
    return JsonResponse({"message": "Hello backend!"})


# TODO: separate into different files all handlers

@api_view(['POST'])
def import_file(request):
    file = request.FILES.get("file")

    if not file:
        logger.debug("import: no file uploaded")
        return JsonResponse({"error": "No file uploaded"}, status=400)

    filename = file.name
    logger.debug(f"import: got file: {filename}")

    result = ""
    if filename.endswith(".txt"):
        result = file.read().decode("utf-8")
    elif filename.endswith(".docx"):
        doc = Document(BytesIO(file.read()))
        for paragraph in doc.paragraphs:
            result += paragraph.text + "\n"
    elif filename.endswith(".pdf"):
        reader = PdfReader(BytesIO(file.read()))
        for page in reader.pages:
            text = page.extract_text()

            # TODO: add paragraph newlines from source pdf
            if text:
                result += text + "\n"

    return JsonResponse({
        "text": result,
    })


@api_view(['POST'])
def export_docx(request):
    text = request.data.get("text", "")

    document = Document()

    lines = text.split("\n")

    # TODO: add font?

    for line in lines:
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename="document.docx"'

    return response